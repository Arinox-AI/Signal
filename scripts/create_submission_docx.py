from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "Signal_Technical_and_Product_Documentation.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5E6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALL_OUT = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "161B22"
USABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_run(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_font(run, **kwargs)
    return run


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_page_field(paragraph):
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instr)
    run._r.append(field_end)
    set_font(run, size=8.5, color=MUTED)


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    style_specs = {
        "Title": (23, NAVY, 0, 4),
        "Subtitle": (13, MUTED, 0, 16),
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in style_specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Small Note" not in [s.name for s in doc.styles]:
        note = doc.styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = doc.styles["Small Note"]
    note.font.name = "Calibri"
    note._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    note._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    note.font.size = Pt(9)
    note.font.color.rgb = RGBColor.from_string(MUTED)
    note.paragraph_format.space_after = Pt(5)
    note.paragraph_format.line_spacing = 1.05

    header = section.header
    paragraph = header.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, after=0, line=1.0)
    style_run(paragraph, "SIGNAL", size=8.5, color=NAVY, bold=True)
    style_run(paragraph, "  |  TECHNICAL & PRODUCT DOCUMENTATION", size=8.5, color=MUTED)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(paragraph, before=0, after=0, line=1.0)
    style_run(paragraph, "Alba Assessment  |  July 2026  |  Page ", size=8.5, color=MUTED)
    add_page_field(paragraph)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=12, after=2, line=1.0)
    style_run(p, "ALBA CORP. CREATIVE, API-INTEGRATED WEB APP", size=9, color=BLUE, bold=True)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_run(p, "Signal", size=27, color=NAVY, bold=True)

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_run(p, "Technical & Product Documentation", size=14, color=MUTED)

    table = doc.add_table(rows=3, cols=2)
    set_table_geometry(table, [2700, 6660])
    rows = [
        ("Project", "Signal — public-data company intelligence"),
        ("Repository", "maluraditya/alba-assessments · /01-web-app"),
        ("Deployment target", "Vercel with 01-web-app configured as the Root Directory"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_cell_shading(row.cells[1], LIGHT_GRAY)
        p1 = row.cells[0].paragraphs[0]
        p2 = row.cells[1].paragraphs[0]
        set_paragraph_spacing(p1, after=0, line=1.05)
        set_paragraph_spacing(p2, after=0, line=1.05)
        style_run(p1, label, size=9.5, color=NAVY, bold=True)
        style_run(p2, value, size=9.5, color=BLACK)

    doc.add_paragraph()
    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [9360])
    set_cell_shading(callout.cell(0, 0), CALL_OUT)
    p = callout.cell(0, 0).paragraphs[0]
    set_paragraph_spacing(p, before=2, after=2, line=1.12)
    style_run(p, "Purpose. ", size=10.5, color=NAVY, bold=True)
    style_run(
        p,
        "This document gives reviewers a complete, source-aware explanation of Signal's product scope, API strategy, architecture, advanced features, and verification approach.",
        size=10.5,
        color=BLACK,
    )


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_body(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6, line=1.1)
    style_run(p, text, size=11, color=BLACK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=4, line=1.15)
    style_run(p, text, size=10.8, color=BLACK)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(p, after=4, line=1.15)
    style_run(p, text, size=10.8, color=BLACK)
    return p


def add_labelled_paragraph(doc, label, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6, line=1.1)
    style_run(p, f"{label}: ", size=11, color=NAVY, bold=True)
    style_run(p, text, size=11, color=BLACK)
    return p


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    for cell, text in zip(header.cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.0)
        style_run(p, text, size=font_size, color=NAVY, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_values):
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.05)
            style_run(p, text, size=font_size, color=BLACK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_cell_shading(table.cell(0, 0), "EEF2F6")
    p = table.cell(0, 0).paragraphs[0]
    set_paragraph_spacing(p, before=2, after=2, line=1.05)
    for index, line in enumerate(lines):
        run = p.add_run(line)
        set_font(run, name="Consolas", size=9.5, color="243447")
        if index < len(lines) - 1:
            run.add_break()
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_document():
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    add_heading(doc, "1. Product overview")
    add_body(doc, "Signal is a public-data company intelligence product. A user enters a company name, legal entity, or official domain and receives one source-aware briefing that combines company identity, official-web evidence, developer activity, current coverage, country context, and an optional AI synthesis.")
    add_body(doc, "The product is intentionally not a generic dashboard. Its interface is an editorial research environment: the landing page explains the trust model before search, while the report presents evidence as a decision brief rather than a collection of disconnected widgets.")
    add_heading(doc, "User problem", 2)
    add_body(doc, "Researching an unfamiliar company normally means moving between search engines, legal-entity records, company websites, GitHub, news feeds, and AI tools. This is slow and increases the risk of confusing similarly named organisations or treating a missing source as a negative signal.")
    add_heading(doc, "Product outcome", 2)
    add_body(doc, "Signal reduces that work to a shareable URL and a clear report. It makes the evidence found, the evidence unavailable or uncertain, and the next verification step explicit.")
    add_heading(doc, "Scope", 2)
    add_labelled_paragraph(doc, "Included", "Global debounced search, multi-source evidence fusion, evidence-constrained Gemini synthesis, source-level loading/empty/error states, and a responsive accessible research experience.")
    add_labelled_paragraph(doc, "Deliberately excluded", "Accounts, saved lists, billing, CRM integrations, private company databases, paid data providers, and invented analytics. These would dilute the search-to-brief experience within the assessment time-box.")

    add_heading(doc, "2. API choices")
    add_body(doc, "Signal uses public, free, and key-optional sources. No credential is exposed to the browser; every upstream request is made by server-only services.")
    add_table(
        doc,
        ["Source", "Role in Signal", "Why this choice", "Failure behaviour"],
        [
            ("GLEIF", "Global legal-entity search and LEI selection", "Broad registered-entity coverage beyond technology companies", "Search continues with alternate identity sources"),
            ("Wikidata + Wikipedia", "Organisation matching and public context", "Useful canonical public-identity path for notable companies", "Report continues with website or legal-entity evidence"),
            ("Official website metadata", "Name, description, founding year, location, industry", "First-party evidence for local companies and startups", "Initials and remaining identity signals remain available"),
            ("Public-web discovery", "Name-to-official-domain fallback", "Resolves companies absent from structured indexes", "Shows an explicit no-match state rather than inventing a result"),
            ("GitHub REST API", "Verified public organisation and repository footprint", "Independent builder signal when genuinely verifiable", "Panel reports no signal or rate-limit state independently"),
            ("Google News RSS", "Current coverage and timeline", "Free broad-coverage alternative to a news trial", "News becomes empty/unavailable without blocking report"),
            ("REST Countries v5", "Country operating context", "Adds dependable country-level context", "Country panel degrades independently"),
            ("Google Gemini", "Executive synthesis and watch item", "Turns supplied evidence into a concise brief", "Deterministic source-derived brief replaces AI output"),
        ],
        [1600, 2550, 2850, 2360],
        font_size=8.4,
    )
    add_heading(doc, "API-specific implementation notes", 2)
    add_heading(doc, "Identity resolution and ambiguity", 3)
    add_body(doc, "Wikidata and GLEIF search in parallel. A selected GLEIF record remains tied to its unique LEI rather than being resolved again by fuzzy name, avoiding similarly named legal entities being mixed together. Official-domain discovery is labelled as website evidence rather than legal-record proof.")
    add_heading(doc, "GitHub matching", 3)
    add_body(doc, "Legal suffixes such as Inc. and Ltd. are removed before trying a bounded, deduplicated candidate set. A matched organisation with zero repositories is treated as an empty development signal, not as misleading zero metrics. GITHUB_TOKEN is optional and only raises the public rate limit.")
    add_heading(doc, "Gemini safeguards", 3)
    add_body(doc, "Gemini receives a constrained normalized evidence bundle. Its JSON output is schema-validated before rendering. Missing credentials, quota errors, invalid model output, and upstream failures use the same source-derived fallback.")

    add_heading(doc, "3. Architecture")
    add_heading(doc, "System overview", 2)
    add_code_block(doc, [
        "Browser",
        "  ├─ Server-rendered landing and report routes",
        "  ├─ Search, motion, retry, and copy-link interactions",
        "  └─ /api/search and /api/company/[query]",
        "                  │",
        "                  ▼",
        "       Next.js backend-for-frontend",
        "                  │",
        "                  ▼",
        " Typed orchestration → GLEIF / Wikidata / Website / GitHub / News / Countries / Gemini",
        "                  │",
        "                  ▼",
        "       Normalized IntelligenceReport",
    ])
    add_heading(doc, "Client/server boundary", 2)
    add_table(
        doc,
        ["Layer", "Responsibility"],
        [
            ("Server-rendered routes", "Resolve reports, stream secondary sections, generate metadata, and keep first meaningful content visible for performance and SEO."),
            ("Route handlers", "Provide GET /api/search for suggestions and GET /api/company/[query] for programmatic access and retry."),
            ("Server-only services", "Call APIs, manage secrets, enforce timeout/retry policy, normalize payloads, and invoke Gemini."),
            ("Client Components", "Handle input state, debouncing, keyboard selection, focus management, motion, copy-link feedback, and retry interaction."),
            ("Report components", "Render typed normalized data only; they do not hold provider response shapes or credentials."),
        ],
        [2450, 6910],
        font_size=9.2,
    )
    add_heading(doc, "Data model and resilience", 2)
    add_body(doc, "Every provider returns a discriminated SourceResult<T>: success, empty, unavailable, or rate-limited. The orchestration service assembles those outcomes into one IntelligenceReport. Secondary calls begin together with Promise.all after primary identity resolution, avoiding serial network waterfalls. Each request uses a timeout, bounded retry, exponential backoff, and Retry-After awareness where relevant.")
    caching_heading = add_heading(doc, "Caching strategy", 2)
    caching_heading.paragraph_format.page_break_before = True
    add_table(
        doc,
        ["Data", "Typical server-side cache lifetime"],
        [
            ("Google News RSS", "10 minutes"),
            ("GitHub repositories / organisation", "15 minutes / 30 minutes"),
            ("Official website metadata", "24 hours"),
            ("Gemini synthesis", "24 hours"),
            ("GLEIF legal-entity result", "6 hours"),
            ("Wikipedia / Wikidata", "24 hours to 7 days"),
            ("REST Countries", "7 days"),
        ],
        [4500, 4860],
        font_size=9.5,
    )
    add_heading(doc, "Deployment architecture", 2)
    add_body(doc, "The repository is a small assessment monorepo. Signal lives in /01-web-app; /02-dashboard and /03-n8n-workflow are reserved for the other tasks. Vercel should import the repository with Root Directory set to 01-web-app. The Signal folder has its own package.json, lockfile, environment template, documentation, and Next.js configuration, so it builds independently.")

    add_heading(doc, "4. Advanced features")
    add_heading(doc, "Backend-for-frontend", 2)
    add_body(doc, "Next.js route handlers and server-only services aggregate upstream calls, keep credentials private, normalize provider responses, cache source results, and translate rate-limit or failure conditions into stable UI states. The browser never needs to understand seven provider response formats.")
    add_heading(doc, "Multi-API data fusion", 2)
    add_body(doc, "One company view fuses selected identity, official web evidence, verified GitHub activity where available, recent coverage, country context, and an evidence-aware executive brief. The resulting timeline and decision brief cannot be provided by any one source alone.")
    add_heading(doc, "Shareable URL-synchronised search", 2)
    add_body(doc, "Search is debounced and server-backed. Suggestions support pointer, touch, and keyboard interaction, then resolve to stable URLs such as /company/vercel. The report can be copied or shared directly, and browser history remains meaningful.")
    add_heading(doc, "Signature motion system", 2)
    add_body(doc, "The landing page uses pointer-responsive parallax, orbiting source nodes, scroll-triggered reveals, a source constellation, a motion marquee, and a spring-smoothed progress rail. Motion is transform/opacity-oriented and respects reduced-motion preferences. It explains evidence flow and hierarchy rather than acting as decoration alone.")

    add_heading(doc, "5. Experience, accessibility, and performance")
    add_heading(doc, "Loading and failure states", 2)
    for item in [
        "Report loading uses section-shaped skeletons and shimmer rather than a generic spinner.",
        "Search has explicit empty and no-match states, with official-domain guidance instead of a fabricated company.",
        "Evidence panels expose empty, unavailable, and rate-limited states without removing the remaining report.",
        "Retry actions are available where a user can reasonably retry a transient failure.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Accessibility", 2)
    for item in [
        "Semantic landmarks and heading hierarchy provide a clear document outline.",
        "Search provides an accessible label, live announcements, keyboard selection, and visible focus states.",
        "Interactive controls are keyboard reachable and have accessible names.",
        "Colour contrast targets WCAG AA; source state is communicated with wording and icons, not colour alone.",
        "Reduced-motion users receive complete content without non-essential animation.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Performance choices", 2)
    for item in [
        "Server Components are the default; client JavaScript is limited to interaction and motion.",
        "Independent upstream calls run in parallel and slower sections stream behind meaningful Suspense fallbacks.",
        "Hero text stays server-visible rather than being delayed behind animation.",
        "Native system typography avoids a render-blocking remote font request.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. Testing and verification")
    add_heading(doc, "Automated release checks", 2)
    add_code_block(doc, [
        "cd 01-web-app",
        "npm run format:check",
        "npm run typecheck",
        "npm run lint",
        "npm run test",
        "npm run build",
    ])
    add_body(doc, "The final local release check passes Prettier formatting, strict TypeScript, zero-warning ESLint, 9 Vitest tests, and the optimized Next.js production build.")
    manual_checks_heading = add_heading(doc, "Manual product checks", 2)
    manual_checks_heading.paragraph_format.page_break_before = True
    add_table(
        doc,
        ["Scenario", "What was verified"],
        [
            ("Technology companies", "Vercel, Stripe, Perplexity, Nintendo, and OpenAI: identity resolution, GitHub activity, coverage, timeline, and adaptive report layout."),
            ("Non-technology companies", "Coca-Cola, Toyota, and McDonald's: useful report composition without public developer activity."),
            ("Local / less-indexed company", "Alba Corporation UAE: official-domain discovery, structured website identity, UAE country context, and source-aware brief."),
            ("No match", "Nonsense query: clear no-match state with official-domain guidance."),
            ("Search interaction", "Seven-result Toyota dropdown: pointer scrolling, keyboard navigation to final item, and mobile-width behaviour."),
            ("Source degradation", "Missing-key and unavailable-source states preserve the rest of the report shell and evidence."),
            ("Responsive layout", "No horizontal overflow at 360, 768, 1024, 1440, and 2560 px."),
        ],
        [2500, 6860],
        font_size=9.1,
    )
    add_heading(doc, "Lighthouse result", 2)
    add_table(
        doc,
        ["Category", "Score"],
        [("Performance", "97"), ("Accessibility", "100"), ("Best Practices", "100"), ("SEO", "100")],
        [6200, 3160],
        font_size=10,
    )
    note = doc.add_paragraph(style="Small Note")
    style_run(note, "Observed mobile metrics: LCP 2.6 s, CLS 0, TBT 10 ms.", size=9, color=MUTED, italic=True)
    add_heading(doc, "Remaining release checks", 2)
    for item in [
        "Deploy through Vercel with Root Directory set to 01-web-app.",
        "Add production environment variables from .env.example.",
        "Set NEXT_PUBLIC_APP_URL to the final deployment URL and redeploy.",
        "Smoke-test /, /company/vercel, /api/company/vercel, /robots.txt, /sitemap.xml, and the social card in a signed-out browser.",
        "Record a short walkthrough showing the search flow, a multi-source result, and one engineering challenge.",
    ]:
        add_numbered(doc, item)

    environment_heading = add_heading(doc, "7. Environment and local setup")
    environment_heading.paragraph_format.page_break_before = True
    add_code_block(doc, ["cd 01-web-app", "npm install", "cp .env.example .env.local", "npm run dev"])
    add_table(
        doc,
        ["Variable", "Required", "Purpose"],
        [
            ("GEMINI_API_KEY", "For AI briefs", "Server-side Gemini access"),
            ("REST_COUNTRIES_API_KEY", "For country context", "REST Countries v5 access"),
            ("GITHUB_TOKEN", "No", "Raises public GitHub API rate limits"),
            ("NEXT_PUBLIC_APP_URL", "In production", "Canonical metadata and sitemap URL"),
            ("NEXT_PUBLIC_SOURCE_URL", "In production", "Public repository link in the header"),
            ("GEMINI_MODEL", "No", "Optional Gemini model override"),
        ],
        [2900, 1750, 4710],
        font_size=9.3,
    )
    note = doc.add_paragraph(style="Small Note")
    style_run(note, "Security note: never commit .env.local or a real credential. Only .env.example is versioned.", size=9, color=MUTED, italic=True)

    add_heading(doc, "8. Repository guide")
    add_table(
        doc,
        ["Location", "Purpose"],
        [
            ("app/", "App Router pages, API routes, metadata routes, loading, and error boundaries"),
            ("components/", "Small landing, report, interaction, motion, and UI components"),
            ("services/", "Server-only provider adapters and orchestration"),
            ("lib/", "Typed contracts, normalisation, formatting, and request utilities"),
            ("docs/", "Product, technical, design, test, deployment, and submission documentation"),
            ("BUILD_LOG.md", "Concise engineering decision journal for the assessment"),
        ],
        [2500, 6860],
        font_size=9.4,
    )
    add_body(doc, "For deeper implementation references, see the PRD, technical specification, UI/UX specification, test plan, deployment guide, and assessment checklist in the docs directory.")

    # Ensure sections begin naturally after the title block and retain the business brief structure.
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Signal — Technical & Product Documentation"
    doc.core_properties.subject = "Alba Assessment submission documentation"
    doc.core_properties.author = "Signal"
    doc.core_properties.comments = "Editable Word document generated from the Signal submission documentation."
    doc.save(OUT)


if __name__ == "__main__":
    build_document()
    print(OUT)
